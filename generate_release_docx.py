from pathlib import Path

from docx import Document
from docx.shared import Inches


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for key, value in rows:
        r = table.add_row().cells
        r[0].text = str(key)
        r[1].text = str(value)
    doc.add_paragraph("")


def add_numbered(doc, items):
    for i, text in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {text}")
    doc.add_paragraph("")


def add_screen(doc, title, image_path):
    doc.add_heading(title, level=3)
    if image_path.exists():
        doc.add_picture(str(image_path), width=Inches(6.8))
    else:
        doc.add_paragraph(f"[Missing screenshot: {image_path.name}]")
    doc.add_paragraph("")


def main():
    base = Path("d:/canal_audit_system")
    out_file = base / "release" / "CAG_Audit_Management_System_Release_Document_v1.0.docx"
    screenshot_dir = base / "release" / "screenshots"

    doc = Document()
    doc.add_heading("CAG Audit Management System", level=0)
    doc.add_paragraph("Production Release Document (Government Submission Copy)")

    add_kv_table(
        doc,
        [
            ("Document ID", "CAG-AMS-REL-1.0"),
            ("Version", "1.0"),
            ("Release Date", "09 April 2026"),
            ("Prepared For", "Comptroller and Auditor General of India (CAG)"),
            ("System", "Audit Management System - Irrigation Audit Wing"),
            ("Prepared By", "Application Engineering Team"),
        ],
    )

    doc.add_heading("1. Document Control", level=1)
    add_kv_table(
        doc,
        [
            ("Document Type", "Production Release Note and Screen Register"),
            ("Classification", "Internal Government Use"),
            ("Target Environment", "Production"),
            ("Application Stack", "Streamlit + Python + PostgreSQL"),
            ("Release Scope", "Operator and Administrator workflows"),
        ],
    )

    doc.add_heading("2. Release Summary", level=1)
    add_numbered(
        doc,
        [
            "Secure authentication and role-based access.",
            "Operator-side estimate creation and section-wise data entry.",
            "Administrator-side review, user creation, and permission management.",
            "Draft and submission lifecycle support with module-wise progress tracking.",
        ],
    )

    doc.add_heading("3. Validation Summary (Release Readiness Evidence)", level=1)
    doc.add_heading("3.1 Functional and Unit Validation", level=2)
    add_numbered(
        doc,
        [
            "Unit and integration test suite executed: 57/57 PASSED.",
            "Authentication, user management, submission creation, and section save flows validated.",
        ],
    )
    doc.add_heading("3.2 Concurrent Load Validation", level=2)
    add_numbered(
        doc,
        [
            "Target requested: 300 users at a time.",
            "Validation command executed: python load_test.py --users 300 --concurrency 300 --module contract_management --fill-scope first --cleanup --fail-on-errors --json-out .tmp\\loadtest_300_fixed.json",
            "Result: 300 success, 0 failure, 100% success rate.",
        ],
    )

    doc.add_heading("4. Deployment and Configuration Notes", level=1)
    doc.add_heading("4.1 Mandatory Runtime Settings", level=2)
    add_numbered(
        doc,
        [
            "DB_HOST, DB_NAME, DB_USER, DB_PASSWORD, DB_PORT",
            "COOKIE_PASSWORD (minimum 16 characters)",
        ],
    )
    doc.add_heading("4.2 Connection Pool Tuning Parameters", level=2)
    add_numbered(
        doc,
        [
            "DB_POOL_MINCONN (default 1)",
            "DB_POOL_MAXCONN (default 80)",
            "DB_POOL_WAIT_TIMEOUT (default 30 seconds)",
            "DB_POOL_WAIT_POLL_INTERVAL (default 0.05 seconds)",
        ],
    )

    doc.add_heading("5. User Roles in Scope", level=1)
    add_numbered(
        doc,
        [
            "Operator: Create estimates, fill sections, save drafts, submit applications.",
            "Administrator: Review applications, create users, revoke/grant access, update module permissions.",
        ],
    )

    doc.add_heading("6. Screen Register With Release Screenshots", level=1)
    screens = [
        ("6.1 Login Page", "01_login_page.png"),
        ("6.2 Operator Dashboard", "02_operator_dashboard.png"),
        ("6.3 New Estimate Dialog (Blank)", "03_new_estimate_dialog.png"),
        ("6.4 New Estimate Dialog (Filled)", "04_new_estimate_filled.png"),
        ("6.5 Contract Management Module - Overview", "05_module_form_overview.png"),
        ("6.6 Tab: Admin Financial Sanction", "06_tab_admin_financial_sanction.png"),
        ("6.7 Tab: Technical Sanction", "07_tab_technical_sanction.png"),
        ("6.8 Tab: Tender Award Contract", "08_tab_tender_award_contract.png"),
        ("6.9 Tab: Contract Master", "09_tab_contract_master.png"),
        ("6.10 Tab: Payments Recoveries", "10_tab_payments_recoveries.png"),
        ("6.11 Tab: Budget Summary", "11_tab_budget_summary.png"),
        ("6.12 Post Logout Screen", "12_login_after_logout.png"),
        ("6.13 Admin Panel: Review Applications", "13_admin_review_applications.png"),
        ("6.14 Admin Panel: Create User", "14_admin_create_user.png"),
        ("6.15 Admin Panel: Manage Users", "15_admin_manage_users.png"),
        ("6.16 Admin Panel: Review Applications (All Users)", "16_admin_review_all_users.png"),
        ("6.17 Operator Estimate Dialog", "17_operator_estimate_dialog.png"),
    ]
    for title, name in screens:
        add_screen(doc, title, screenshot_dir / name)

    doc.add_heading("7. Operational Runbook", level=1)
    add_numbered(
        doc,
        [
            "Start command: streamlit run app.py",
            "Verify login, operator data save, and admin review pages.",
            "Use previous release package and DB backup for rollback when required.",
        ],
    )

    doc.add_heading("8. Security and Governance Checklist", level=1)
    add_numbered(
        doc,
        [
            "Role-based access controls validated.",
            "Password hashing and verification enabled.",
            "Cookie secret policy enforced.",
            "Connection pool controls added for concurrency resilience.",
            "Release evidence includes functional and load validation outputs.",
        ],
    )

    doc.add_heading("9. Sign-off Section", level=1)
    add_kv_table(
        doc,
        [
            ("Project Owner", ""),
            ("Application Administrator", ""),
            ("Information Security Officer", ""),
            ("Deployment Authority", ""),
        ],
    )

    doc.add_heading("10. Release Package Contents", level=1)
    add_numbered(
        doc,
        [
            "Source application files (app.py, auth.py, crud.py, supporting assets)",
            "Release documentation files in release/",
            "Screen evidence in release/screenshots/",
            "Load validation evidence in .tmp/loadtest_300_fixed.json",
        ],
    )

    out_file.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_file))
    print(str(out_file))


if __name__ == "__main__":
    main()
